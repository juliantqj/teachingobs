#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri A1 ug 25 11:28:42 2025

@author: komarjo1
"""

import whisper
from pathlib import Path
import csv
import time
from datetime import timedelta
import torch
from docx import Document

start_time = time.time()


## If ffmpeg not found
# os
#os.environ["PATH"] += os.pathsep + "/opt/homebrew/Cellar"


# ---- CONFIG ----
input_folder = Path("input")
export_folder = Path("export")
model_size = 'large-v3'  # Choose from: tiny, base, small, medium, large, large-v3, etc.
valid_extensions = [".mp3", ".m4a", ".wav"]

# ---- SETUP EXPORT FOLDER ----
export_folder.mkdir(parents=True, exist_ok=True)

# ---- DEVICE DETECTION ----
if torch.cuda.is_available():
    device = "cuda"
elif torch.backends.mps.is_available():
    device = "cpu"
else:
    device = "cpu"

print(f"Using device: {device}")

# ---- LOAD MODEL ----
model = whisper.load_model(model_size, device=device)

# ---- PROCESS ALL AUDIO FILES ----
audio_files = [f for f in input_folder.iterdir() if f.suffix.lower() in valid_extensions]

if not audio_files:
    print("No audio files found in the input folder.")
else:
    for audio_path in audio_files:
        print(f"\nProcessing file: {audio_path.name}")
        base_name = audio_path.stem
        output_docx_filename = export_folder / f"{base_name}_transcript.docx"
        output_csv_filename = export_folder / f"{base_name}_segments.csv"
        output_txt_filename = export_folder / f"{base_name}_transcript.txt"

        # ---- TRANSCRIBE ----
        result = model.transcribe(str(audio_path), language='en', verbose=True)

        # ---- EXPORT WORD DOCUMENT (.docx) ----
        document = Document()
        document.add_heading(f"Transcript: {audio_path.name}", level=1)
        document.add_paragraph(result['text'])
        document.save(output_docx_filename)
        print(f"Transcript saved to Word document: {output_docx_filename.name}")

        # ---- EXPORT PLAIN TEXT (.txt) ----
        with open(output_txt_filename, "w", encoding="utf-8") as txt_file:
            txt_file.write(result['text'])
        print(f"Transcript saved to plain text file: {output_txt_filename.name}")

        # ---- EXPORT CSV WITH TIMESTAMPS ----
        segments = result.get("segments", [])
        if segments:
            with open(output_csv_filename, "w", newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(["segment", "start", "end", "text"])
                for i, seg in enumerate(segments):
                    writer.writerow([
                        i,
                        round(seg["start"], 2),
                        round(seg["end"], 2),
                        seg["text"].strip()
                    ])
            print(f"CSV with segments saved to: {output_csv_filename.name}")
        else:
            print("No segments found in transcription result.")

# ---- TIMER ----
elapsed_time_secs = time.time() - start_time
print("\nTotal execution time:", timedelta(seconds=round(elapsed_time_secs)))
