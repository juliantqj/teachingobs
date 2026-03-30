"""
ACAI — audio_analysis.py
Thin wrapper around audio_transcript_JK.py logic.
Accepts a single audio file path + output directory.
Core transcription logic is UNCHANGED from audio_transcript_JK.py.
"""

import csv
import time
from pathlib import Path
from datetime import timedelta



def run_transcription(audio_path: str, output_dir: str,
                      model_size: str = "large-v3",
                      progress_callback=None) -> dict:
    """
    Transcribe a single audio file.
    Returns dict with paths to all output files.
    progress_callback: optional callable(message: str)
    """
    import whisper
    from docx import Document
    import torch

    audio_path  = Path(audio_path)
    output_dir  = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    base_name           = audio_path.stem
    output_docx         = output_dir / f"{base_name}_transcript.docx"
    output_csv          = output_dir / f"{base_name}_segments.csv"
    output_txt          = output_dir / f"{base_name}_transcript.txt"

    # ---- DEVICE DETECTION (unchanged from audio_transcript_JK.py) ----
    if torch.cuda.is_available():
        device = "cuda"
    elif torch.backends.mps.is_available():
        device = "cpu"
    else:
        device = "cpu"

    if progress_callback:
        progress_callback(f"Using device: {device}")
        progress_callback(f"Loading Whisper {model_size}...")

    # ---- LOAD MODEL (unchanged) ----
    model = whisper.load_model(model_size, device=device)

    if progress_callback:
        progress_callback(f"Transcribing: {audio_path.name}")

    # ---- TRANSCRIBE (unchanged) ----
    result = model.transcribe(str(audio_path), language="en", verbose=False)

    # ---- EXPORT WORD DOCUMENT (unchanged) ----
    document = Document()
    document.add_heading(f"Transcript: {audio_path.name}", level=1)
    document.add_paragraph(result["text"])
    document.save(output_docx)

    # ---- EXPORT PLAIN TEXT (unchanged) ----
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write(result["text"])

    # ---- EXPORT CSV WITH TIMESTAMPS (unchanged) ----
    segments = result.get("segments", [])
    with open(output_csv, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(["segment", "start", "end", "text"])
        for i, seg in enumerate(segments):
            writer.writerow([
                i,
                round(seg["start"], 2),
                round(seg["end"], 2),
                seg["text"].strip()
            ])

    if progress_callback:
        progress_callback(f"Transcription complete: {len(segments)} segments")

    return {
        "transcript_txt":  str(output_txt),
        "transcript_docx": str(output_docx),
        "segments_csv":    str(output_csv),
        "full_text":       result["text"],
        "segments":        segments,
    }
